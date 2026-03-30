#!/usr/bin/env python3
"""Generate a polished reference.docx for Pandoc conversions.

Creates common Word styles (Normal, Title, Heading 1-3, Code, Blockquote)
so `pandoc --reference-doc=reference.docx` produces a professional-looking output.
"""
import argparse
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def set_style_font(style, name='Calibri', size=11, bold=False, italic=False):
    font = style.font
    font.name = name
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('-o', '--out', default='reference.docx', help='Output reference docx')
    args = parser.parse_args()

    doc = Document()

    # Title page sample (will be ignored by pandoc but helps Word styles)
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('Project Readme')
    run.bold = True
    run.font.size = Pt(28)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = subtitle.add_run('Generated reference template')
    run2.font.size = Pt(12)
    run2.italic = True

    doc.add_page_break()

    # Configure styles
    styles = doc.styles

    # Normal
    normal = styles['Normal']
    set_style_font(normal, name='Calibri', size=11, bold=False)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    # Title
    if 'Title' in styles:
        t = styles['Title']
        set_style_font(t, name='Calibri Light', size=26, bold=True)

    # Headings
    if 'Heading 1' in styles:
        h1 = styles['Heading 1']
        set_style_font(h1, name='Calibri', size=20, bold=True)
        h1.paragraph_format.space_before = Pt(12)
        h1.paragraph_format.space_after = Pt(6)

    if 'Heading 2' in styles:
        h2 = styles['Heading 2']
        set_style_font(h2, name='Calibri', size=16, bold=True)
        h2.paragraph_format.space_before = Pt(10)
        h2.paragraph_format.space_after = Pt(4)

    if 'Heading 3' in styles:
        h3 = styles['Heading 3']
        set_style_font(h3, name='Calibri', size=13, bold=True)
        h3.paragraph_format.space_before = Pt(8)
        h3.paragraph_format.space_after = Pt(4)

    # Code style (create a character style if available)
    try:
        from docx.enum.style import WD_STYLE_TYPE
        code_style = styles.add_style('Code', WD_STYLE_TYPE.CHARACTER)
        set_style_font(code_style, name='Consolas', size=10)
    except Exception:
        pass

    # Blockquote: use Intense Quote or create a paragraph style
    if 'Intense Quote' in styles:
        iq = styles['Intense Quote']
        set_style_font(iq, name='Calibri', size=11, italic=True)

    # Add a sample heading and paragraph to ensure styles are embedded
    doc.add_heading('Sample Heading 1', level=1)
    doc.add_paragraph('This is a sample paragraph to embed the Normal style and spacing.')
    doc.add_heading('Sample Heading 2', level=2)
    doc.add_paragraph('Another paragraph under heading 2. Lists, code blocks and tables will inherit from these base styles.')

    # Footer note
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.text = 'Generated reference template — adjust in Word if desired.'
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if p.runs:
        p.runs[0].font.size = Pt(9)

    doc.save(args.out)
    print(f'Wrote {args.out}')


if __name__ == '__main__':
    main()
